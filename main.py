from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import zipfile
import io

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    uploaded_file = request.files['file']
    if uploaded_file.filename != '' and uploaded_file.filename.endswith('.docx'):
        # שמירת הקובץ שהועלה
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_file.filename)
        uploaded_file.save(file_path)

        # חילוץ הנספחים מהקובץ
        appendices = extract_appendices_from_word(file_path)

        if not appendices:
            return "לא נמצאו נספחים בקובץ. ווודא שהקובץ מכיל טקסט כמו 'נספח 1', 'נספח 2' וכו'."

        # יצירת שני הקבצים
        table_file = create_appendices_table(appendices)
        covers_file = create_appendices_covers(appendices)

        # יצירת קובץ ZIP עם שני הקבצים
        zip_path = create_zip_file(table_file, covers_file)

        return send_file(zip_path, as_attachment=True, download_name='נספחים.zip')

    return "יש להעלות קובץ Word (.docx)"

def extract_appendices_from_word(file_path):
    """
    פונקציה שקוראת קובץ Word ומחלצת את כל הנספחים
    מחזירה רשימה של tuples: [(מספר_נספח, שם_נספח), ...]
    """
    doc = Document(file_path)
    appendices = []

    # דפוסים אפשריים לזיהוי נספחים
    patterns = [
        r'נספח\s+(\d+)\s*[-–—:]\s*(.+)',  # נספח 1 - שם הנספח
        r'נספח\s+(\d+)\s+(.+)',            # נספח 1 שם הנספח
        r'נספח\s+מספר\s+(\d+)\s*[-–—:]\s*(.+)',  # נספח מספר 1 - שם
        r'נספח\s+מס[\'׳]\s*(\d+)\s*[-–—:]\s*(.+)',  # נספח מס' 1 - שם
    ]

    found_appendices = {}  # מילון למניעת כפילויות

    # עבור כל פסקה במסמך
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # נסה כל דפוס
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                appendix_num = int(match.group(1))
                appendix_name = match.group(2).strip()

                # הסרת תווים מיותרים מהתחלה ומסוף השם
                appendix_name = re.sub(r'^[-–—:\s]+', '', appendix_name)
                appendix_name = re.sub(r'[.,;:\s]+$', '', appendix_name)

                # הוסף רק אם עדיין לא נמצא או אם השם ארוך יותר
                if appendix_num not in found_appendices or len(appendix_name) > len(found_appendices[appendix_num]):
                    found_appendices[appendix_num] = appendix_name

    # המרה לרשימה ממוינת
    appendices = [(num, name) for num, name in sorted(found_appendices.items())]

    return appendices

def create_appendices_table(appendices):
    """
    יוצר מסמך Word עם טבלה של הנספחים
    """
    doc = Document()

    # הוספת כותרת
    heading = doc.add_heading('רשימת נספחים', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # יצירת טבלה עם 2 עמודות + שורת כותרת
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # שורת כותרת
    header_cells = table.rows[0].cells
    header_cells[0].text = 'מספר נספח'
    header_cells[1].text = 'שם הנספח'

    # עיצוב כותרת
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(14)

    # הוספת שורות עבור כל נספח
    for num, name in appendices:
        row_cells = table.add_row().cells
        row_cells[0].text = f'נספח {num}'
        row_cells[1].text = name

        # יישור לימין
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(12)

    # שמירת הקובץ
    output_path = os.path.join(OUTPUT_FOLDER, 'טבלת_נספחים.docx')
    doc.save(output_path)
    return output_path

def create_appendices_covers(appendices):
    """
    יוצר מסמך Word עם שער לכל נספח (עמוד נפרד)
    """
    doc = Document()

    for i, (num, name) in enumerate(appendices):
        # הוספת מעבר עמוד (חוץ מהעמוד הראשון)
        if i > 0:
            doc.add_page_break()

        # הוספת 5 שורות ריקות מלמעלה
        for _ in range(5):
            doc.add_paragraph()

        # הוספת מספר הנספח (במרכז, גדול ומודגש)
        appendix_num_para = doc.add_paragraph()
        appendix_num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = appendix_num_para.add_run(f'נספח {num}')
        run.bold = True
        run.font.size = Pt(24)

        # הוספת שורה ריקה
        doc.add_paragraph()

        # הוספת שם הנספח (במרכז, בינוני)
        appendix_name_para = doc.add_paragraph()
        appendix_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = appendix_name_para.add_run(name)
        run.font.size = Pt(18)

    # שמירת הקובץ
    output_path = os.path.join(OUTPUT_FOLDER, 'שערי_נספחים.docx')
    doc.save(output_path)
    return output_path

def create_cover_page(appendices, output_path):
    """יצירת דף שער נפרד לכל נספח — כל נספח בעמוד משלו"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(2)
        section.bottom_margin = Inches(2)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def add_centered_paragraph(text, font_size, bold=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = 'David'
        return p

    for i, (num, desc) in enumerate(appendices):
        if i > 0:
            p_break = doc.add_paragraph()
            run_break = p_break.add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run_break._element.append(br)

        for _ in range(6):
            doc.add_paragraph()

        add_centered_paragraph(f'נספח {num}', font_size=32, bold=True, space_before=0, space_after=12)
        add_centered_paragraph('─' * 40, font_size=12, bold=False, space_before=6, space_after=6)
        add_centered_paragraph(desc, font_size=18, bold=False, space_before=12, space_after=0)

    doc.save(output_path)
    print(f'✓ נשמר: {output_path}')


def create_zip_file(table_file, covers_file):
    """
    יוצר קובץ ZIP עם שני הקבצים
    """
    zip_path = os.path.join(OUTPUT_FOLDER, 'נספחים.zip')

    with zipfile.ZipFile(zip_path, 'w') as zipf:
        zipf.write(table_file, os.path.basename(table_file))
        zipf.write(covers_file, os.path.basename(covers_file))

    return zip_path

if __name__ == '__main__':
    app.run(debug=True)
